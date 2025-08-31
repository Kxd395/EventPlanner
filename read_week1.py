#!/usr/bin/env python3
"""
Simple Document Reader for Week 1 Coursework
"""

from docx import Document
import os

def read_week1_document():
    file_path = "/Users/VScode_Projects/review/Week 1 boards and cource work.docx"
    
    try:
        print("Reading Week 1 Coursework Document...")
        print("=" * 50)
        
        doc = Document(file_path)
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"\nParagraph {i+1}:")
                print(paragraph.text)
                print("-" * 30)
        
        # Check for tables
        if doc.tables:
            print(f"\nFound {len(doc.tables)} table(s):")
            for i, table in enumerate(doc.tables):
                print(f"\nTable {i+1}:")
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        print("  | ".join(row_data))
        
        print("\n" + "=" * 50)
        print("Document reading completed!")
        
    except Exception as e:
        print(f"Error reading document: {e}")

if __name__ == "__main__":
    read_week1_document()
