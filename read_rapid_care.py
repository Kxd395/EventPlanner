#!/usr/bin/env python3
"""
Read Rapid Care Inc Document
"""

from docx import Document
import os

def read_rapid_care_document():
    file_path = "/Users/VScode_Projects/review/Rapid Care inc.docx"
    
    try:
        print("Reading Rapid Care Inc. Project Document...")
        print("=" * 50)
        
        doc = Document(file_path)
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"\nParagraph {i+1}:")
                print(paragraph.text)
                print("-" * 30)
        
        # Check for tables
        if doc.tables:
            print(f"\nFound {len(doc.tables)} table(s):")
            for i, table in enumerate(doc.tables):
                print(f"\nTable {i+1}:")
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        print("  | ".join(row_data))
        
        print("\n" + "=" * 50)
        print("Document reading completed!")
        
    except Exception as e:
        print(f"Error reading document: {e}")

if __name__ == "__main__":
    read_rapid_care_document()
